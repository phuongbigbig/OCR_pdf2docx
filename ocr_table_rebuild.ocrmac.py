import os
import sys
import subprocess
import argparse
import time

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Error: Missing python-docx. Run: pip install python-docx")
    sys.exit(1)

try:
    import fitz  # PyMuPDF
except ImportError:
    print("❌ Error: Missing PyMuPDF. Run: pip install pymupdf")
    sys.exit(1)

try:
    import pandas as pd
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import (
        PdfPipelineOptions,
        OcrMacOptions,
        TableFormerMode,
    )
    from docling_core.types.doc import (
        TableItem,
        TextItem,
        SectionHeaderItem,
        ListItem,
        PictureItem,
    )
except ImportError:
    print(
        "❌ Error: Missing core packages. Run:\n"
        '   pip install "docling[ocrmac]" docling-core python-docx pymupdf pandas'
    )
    sys.exit(1)


def open_file_locally(file_path):
    """Opens the converted file using the system default application."""
    try:
        if os.name == "nt":
            os.startfile(file_path)
        elif sys.platform == "darwin":
            subprocess.call(("open", file_path))
        else:
            subprocess.call(("xdg-open", file_path))
    except Exception:
        print("⚠️ Conversion complete, but could not open file automatically.")


def build_docx_table(doc, table_item, docling_doc):
    """Translates a Docling TableItem into a native Word table via DataFrame export."""
    try:
        table_df = table_item.export_to_dataframe(doc=docling_doc)
    except Exception as e:
        print(f"   ⚠️ Could not export table to DataFrame: {e}")
        try:
            table_df = table_item.export_to_dataframe()
        except Exception:
            print("   ⚠️ Table export failed entirely. Skipping this table.")
            return

    if table_df.empty:
        return

    num_data_rows = len(table_df)
    num_cols = len(table_df.columns)

    if num_cols == 0:
        return

    has_header = not all(isinstance(c, int) for c in table_df.columns)

    total_rows = num_data_rows + (1 if has_header else 0)
    word_table = doc.add_table(rows=total_rows, cols=num_cols)
    word_table.style = "Table Grid"

    row_offset = 0

    if has_header:
        for col_idx, col_name in enumerate(table_df.columns):
            cell = word_table.cell(0, col_idx)
            cell_text = _clean_cell_value(str(col_name))
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        row_offset = 1

    for row_idx in range(num_data_rows):
        for col_idx in range(num_cols):
            cell_value = table_df.iloc[row_idx, col_idx]
            cell_text = _clean_cell_value(str(cell_value))
            word_table.cell(row_idx + row_offset, col_idx).text = cell_text

    doc.add_paragraph()


def _clean_cell_value(value):
    """Cleans cell text values from DataFrame export."""
    if value in ("None", "nan", "NaN", "<NA>", ""):
        return ""
    return value.strip()


def parse_pages_to_range_tuple(pages_str, total_pages):
    """Parses a string like '2-5' or '3' into a 1-based (start, end) tuple for Docling."""
    if not pages_str:
        return None

    try:
        if "-" in pages_str:
            start, end = map(int, pages_str.replace(" ", "").split("-"))
        else:
            start = int(pages_str.strip())
            end = start

        start = max(1, min(start, total_pages))
        end = max(start, min(end, total_pages))
        return (start, end)
    except ValueError:
        print("⚠️ Warning: Invalid page range format. Processing all pages instead.")
        return None


OCRMAC_LANG_MAP = {
    "en": "en-US",
    "fr": "fr-FR",
    "it": "it-IT",
    "de": "de-DE",
    "es": "es-ES",
    "pt": "pt-BR",
    "zh": "zh-Hans",
    "zh-hans": "zh-Hans",
    "zh-hant": "zh-Hant",
    "ko": "ko-KR",
    "ja": "ja-JP",
    "ru": "ru-RU",
    "uk": "uk-UA",
    "th": "th-TH",
    "vi": "vi-VT",
    "ar": "ar-SA",
    "tr": "tr-TR",
    "id": "id-ID",
    "cs": "cs-CZ",
    "da": "da-DK",
    "nl": "nl-NL",
    "no": "no-NO",
    "ms": "ms-MY",
    "pl": "pl-PL",
    "ro": "ro-RO",
    "sv": "sv-SE",
}


def _to_ocrmac_langs(languages):
    """Converts short language codes (en, vi, fr) to OcrMac locale codes (en-US, vi-VT, fr-FR)."""
    resolved = []
    for lang in languages:
        lang_lower = lang.lower()
        if lang_lower in OCRMAC_LANG_MAP:
            resolved.append(OCRMAC_LANG_MAP[lang_lower])
        elif lang in OCRMAC_LANG_MAP.values():
            # Already a valid OcrMac locale code (e.g., "en-US" passed directly)
            resolved.append(lang)
        else:
            print(f"⚠️ Unknown language code '{lang}', skipping. "
                  f"Valid short codes: {', '.join(sorted(OCRMAC_LANG_MAP.keys()))}")
    return resolved if resolved else ["en-US"]


def build_converter(languages, force_full_page_ocr=False):
    """Builds a DocumentConverter with macOS native OCR and accurate table structure."""
    pipeline_options = PdfPipelineOptions()

    # Use macOS native Vision framework OCR
    pipeline_options.do_ocr = True
    pipeline_options.ocr_options = OcrMacOptions()

    # Convert short codes (en, vi) to OcrMac locale format (en-US, vi-VT)
    if languages:
        ocrmac_langs = _to_ocrmac_langs(languages)
        pipeline_options.ocr_options.lang = ocrmac_langs
        print(f"   🌐 OcrMac languages resolved: {languages} → {ocrmac_langs}")

    if force_full_page_ocr:
        pipeline_options.ocr_options.force_full_page_ocr = True

    # Use accurate table structure model for best results
    pipeline_options.do_table_structure = True
    pipeline_options.table_structure_options.do_cell_matching = True
    pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )
    return converter


def local_table_reconstruction(
    pdf_input,
    docx_output=None,
    languages=None,
    pages_filter=None,
    auto_open=True,
    force_full_page_ocr=False,
):
    """Processes document structures locally and builds native tables inside Word."""
    if not os.path.exists(pdf_input):
        print(f"❌ Error: File '{pdf_input}' does not exist.")
        sys.exit(1)

    if not docx_output:
        docx_output = os.path.splitext(pdf_input)[0] + "_structured_ocrmac.docx"

    if not languages:
        languages = ["en", "vi"]

    # Get total page count via PyMuPDF
    try:
        pdf_doc = fitz.open(pdf_input)
        total_pages = len(pdf_doc)
        pdf_doc.close()
    except Exception as e:
        print(f"❌ Error reading file structure: {e}")
        sys.exit(1)

    page_range_tuple = parse_pages_to_range_tuple(pages_filter, total_pages)

    print(f"\n📂 Source PDF: {os.path.abspath(pdf_input)}")
    print(f"💾 Target Word: {os.path.abspath(docx_output)}")
    print(f"🌐 OCR Engine: macOS Vision (OcrMac)")
    print(f"🌐 OCR Languages: {languages}")
    print(f"🔬 Table Model: ACCURATE (high fidelity)")
    print(f"📷 Full-page OCR: {'Yes (scanned mode)' if force_full_page_ocr else 'No (auto-detect)'}")
    if page_range_tuple:
        print(
            f"📄 Target Page Range: Pages {page_range_tuple[0]} to "
            f"{page_range_tuple[1]} (out of {total_pages} total pages)"
        )
    else:
        print(f"📄 Target Page Range: All Pages (1-{total_pages})")

    # ── Benchmark: total wall-clock ──
    total_start = time.perf_counter()

    print("🔄 Initializing Docling layout analyzer engine (OcrMac)...")

    try:
        # ── Benchmark: converter init ──
        init_start = time.perf_counter()
        converter = build_converter(languages, force_full_page_ocr)
        init_elapsed = time.perf_counter() - init_start
        print(f"   ⏱️  Converter initialized in {init_elapsed:.2f}s")

        print("🧠 Running document parsing neural networks (extracting text & table frames)...")

        convert_kwargs = {}
        if page_range_tuple:
            convert_kwargs["page_range"] = page_range_tuple

        # ── Benchmark: conversion ──
        convert_start = time.perf_counter()
        conversion_result = converter.convert(pdf_input, **convert_kwargs)
        convert_elapsed = time.perf_counter() - convert_start
        docling_doc = conversion_result.document
        print(f"   ⏱️  PDF parsed in {convert_elapsed:.2f}s")

        doc = Document()

        print("📈 Reconstructing formatting trees sequentially into Word output...")

        # ── Benchmark: docx reconstruction ──
        recon_start = time.perf_counter()

        table_count = 0
        text_count = 0

        for item, level in docling_doc.iterate_items():

            if isinstance(item, TableItem):
                table_count += 1
                print(f"   📊 Found table #{table_count}. Mapping cells to Word grid...")
                build_docx_table(doc, item, docling_doc)

            elif isinstance(item, SectionHeaderItem):
                heading_text = item.text.strip()
                if heading_text:
                    heading_level = getattr(item, "level", level)
                    heading_level = max(1, min(heading_level, 9))
                    doc.add_heading(heading_text, level=heading_level)

            elif isinstance(item, ListItem):
                list_text = item.text.strip()
                if list_text:
                    try:
                        doc.add_paragraph(list_text, style="List Bullet")
                    except KeyError:
                        doc.add_paragraph(f"• {list_text}")

            elif isinstance(item, TextItem):
                text_content = item.text.strip()
                if text_content:
                    doc.add_paragraph(text_content)
                    text_count += 1

            elif isinstance(item, PictureItem):
                try:
                    pil_image = item.get_image(docling_doc)
                    if pil_image is not None:
                        import io
                        img_buffer = io.BytesIO()
                        pil_image.save(img_buffer, format="PNG")
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(5.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print("   🖼️ Inserted picture into document.")
                except Exception:
                    caption = getattr(item, "caption", None)
                    if caption:
                        doc.add_paragraph(f"[Figure: {caption}]")

        recon_elapsed = time.perf_counter() - recon_start

        doc.save(docx_output)
        total_elapsed = time.perf_counter() - total_start

        # ── Benchmark summary ──
        print("\n" + "=" * 55)
        print("⏱️  BENCHMARK SUMMARY (OcrMac / macOS Vision)")
        print("=" * 55)
        print(f"   Converter init : {init_elapsed:>8.2f}s")
        print(f"   PDF parsing    : {convert_elapsed:>8.2f}s")
        print(f"   DOCX rebuild   : {recon_elapsed:>8.2f}s")
        print(f"   ─────────────────────────────")
        print(f"   Total wall-clock: {total_elapsed:>7.2f}s")
        print(f"   Pages processed : {page_range_tuple[1] - page_range_tuple[0] + 1 if page_range_tuple else total_pages}")
        if page_range_tuple:
            pages_done = page_range_tuple[1] - page_range_tuple[0] + 1
        else:
            pages_done = total_pages
        print(f"   Speed           : {convert_elapsed / max(pages_done, 1):.2f}s per page (parsing only)")
        print(f"   Tables found    : {table_count}")
        print(f"   Text blocks     : {text_count}")
        print("=" * 55)

        print(
            f"\n🎉 Success! Reconstructed {text_count} text blocks and "
            f"{table_count} tables with fully editable cell text."
        )

        if auto_open:
            open_file_locally(docx_output)

    except Exception as e:
        print(f"\n❌ Layout Analyzer Pipeline Error: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    program_desc = (
        "===============================================================\n"
        "⚡   LOCAL PDF TO WORD CONVERTER  —  macOS Vision OCR Engine  ⚡\n"
        "===============================================================\n"
        "Converts complex scanned or flat PDFs into editable Word files offline.\n"
        "Uses macOS native Vision framework for OCR + Docling neural table models.\n"
        "Includes built-in benchmarking timers for speed comparison."
    )

    epilog_examples = (
        "💡 USAGE EXAMPLES:\n"
        "  1. Standard conversion (Default English + Vietnamese, all pages):\n"
        "     python3 ocr_table_rebuild_v2_ocrmac.py -i internal_report.pdf\n\n"
        "  2. Convert specific page ranges:\n"
        "     python3 ocr_table_rebuild_v2_ocrmac.py -i data.pdf -p 3-7\n\n"
        "  3. Convert with custom language hints:\n"
        "     python3 ocr_table_rebuild_v2_ocrmac.py -i invoice.pdf -l fr en -p 1\n\n"
        "  4. Force full-page OCR for scanned documents:\n"
        "     python3 ocr_table_rebuild_v2_ocrmac.py -i scanned.pdf --scanned\n"
    )

    parser = argparse.ArgumentParser(
        description=program_desc,
        epilog=epilog_examples,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument(
        "-i", "--input",
        help="Path to the source PDF document that you want to convert.",
    )
    parser.add_argument(
        "-o", "--output",
        help="Optional: Custom path/filename for the generated Word (.docx) file.",
    )
    parser.add_argument(
        "-l", "--lang",
        nargs="+",
        metavar="CODE",
        help="Optional: Space-separated language codes for OCR hints (e.g., '-l vi en'). "
             "Short codes (en, vi, fr) are auto-converted to OcrMac locale format.",
    )
    parser.add_argument(
        "-p", "--pages",
        metavar="RANGE",
        help="Optional: Page boundary range to parse. Examples: '2-5' or '4'.",
    )
    parser.add_argument(
        "--scanned",
        action="store_true",
        default=False,
        help="Force full-page OCR on every page (use for scanned/image-based PDFs).",
    )
    parser.add_argument(
        "--no-open",
        action="store_true",
        default=False,
        help="Do not automatically open the output file after conversion.",
    )

    args = parser.parse_args()

    if not args.input:
        print("=== Local Layout-Aware PDF Table Reconstructor (OcrMac) ===")
        user_in = input("Enter PDF file path: ").strip().strip("\"'")
        user_pages = input(
            "Enter specific page range (e.g., '2-5' or '3', blank for all): "
        ).strip()
        user_pages = user_pages if user_pages else None

        print(
            "\n💡 Language Codes: vi (Vietnamese), en (English), "
            "fr (French), ja (Japanese), zh (Chinese Simplified),\n"
            "   ko (Korean), de (German), es (Spanish), th (Thai), "
            "ru (Russian), pt (Portuguese)"
        )
        lang_input = (
            input("Enter language codes space-separated (Leave blank for 'en vi'): ")
            .strip()
            .lower()
        )
        user_langs = lang_input.split() if lang_input else ["en", "vi"]

        scanned_input = (
            input("Is this a scanned/image-based PDF? (y/N): ").strip().lower()
        )
        force_scanned = scanned_input in ("y", "yes")

        if user_in:
            local_table_reconstruction(
                user_in,
                args.output,
                languages=user_langs,
                pages_filter=user_pages,
                force_full_page_ocr=force_scanned,
            )
    else:
        local_table_reconstruction(
            args.input,
            args.output,
            languages=args.lang,
            pages_filter=args.pages,
            auto_open=not args.no_open,
            force_full_page_ocr=args.scanned,
        )
