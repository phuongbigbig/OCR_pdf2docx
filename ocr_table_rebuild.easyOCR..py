import os
import sys
import subprocess
import argparse

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Error: Missing python-docx. Run: pip install python-docx")
    sys.exit(1)

try:
    import fitz  # PyMuPDF
except ImportError:
    print("❌ Error: Missing PyMuPDF. Run: pip install pymupdf")
    sys.exit(1)

try:
    import pandas as pd
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import (
        PdfPipelineOptions,
        EasyOcrOptions,
        TableFormerMode,
    )
    from docling_core.types.doc import (
        TableItem,
        TextItem,
        SectionHeaderItem,
        ListItem,
        PictureItem,
    )
except ImportError:
    print(
        "❌ Error: Missing core packages. Run:\n"
        "   pip install docling docling-core python-docx pymupdf pandas"
    )
    sys.exit(1)


def open_file_locally(file_path):
    """Opens the converted file using the system default application."""
    try:
        if os.name == "nt":
            os.startfile(file_path)
        elif sys.platform == "darwin":
            subprocess.call(("open", file_path))
        else:
            subprocess.call(("xdg-open", file_path))
    except Exception:
        print("⚠️ Conversion complete, but could not open file automatically.")


def build_docx_table(doc, table_item, docling_doc):
    """Translates a Docling TableItem into a native Word table via DataFrame export."""
    try:
        table_df = table_item.export_to_dataframe(doc=docling_doc)
    except Exception as e:
        print(f"   ⚠️ Could not export table to DataFrame: {e}")
        # Fallback: try without doc argument (older docling-core)
        try:
            table_df = table_item.export_to_dataframe()
        except Exception:
            print("   ⚠️ Table export failed entirely. Skipping this table.")
            return

    if table_df.empty:
        return

    num_data_rows = len(table_df)
    num_cols = len(table_df.columns)

    if num_cols == 0:
        return

    # Check if columns are meaningful headers or just integer indices
    has_header = not all(isinstance(c, int) for c in table_df.columns)

    total_rows = num_data_rows + (1 if has_header else 0)
    word_table = doc.add_table(rows=total_rows, cols=num_cols)
    word_table.style = "Table Grid"

    row_offset = 0

    # Write header row if columns have meaningful names
    if has_header:
        for col_idx, col_name in enumerate(table_df.columns):
            cell = word_table.cell(0, col_idx)
            cell_text = _clean_cell_value(str(col_name))
            cell.text = cell_text
            # Bold the header row
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        row_offset = 1

    # Write data rows
    for row_idx in range(num_data_rows):
        for col_idx in range(num_cols):
            cell_value = table_df.iloc[row_idx, col_idx]
            cell_text = _clean_cell_value(str(cell_value))
            word_table.cell(row_idx + row_offset, col_idx).text = cell_text

    doc.add_paragraph()  # spacer after table


def _clean_cell_value(value):
    """Cleans cell text values from DataFrame export."""
    if value in ("None", "nan", "NaN", "<NA>", ""):
        return ""
    return value.strip()


def parse_pages_to_range_tuple(pages_str, total_pages):
    """Parses a string like '2-5' or '3' into a 1-based (start, end) tuple for Docling."""
    if not pages_str:
        return None

    try:
        if "-" in pages_str:
            start, end = map(int, pages_str.replace(" ", "").split("-"))
        else:
            start = int(pages_str.strip())
            end = start

        start = max(1, min(start, total_pages))
        end = max(start, min(end, total_pages))
        return (start, end)
    except ValueError:
        print("⚠️ Warning: Invalid page range format. Processing all pages instead.")
        return None


def build_converter(languages, force_full_page_ocr=False):
    """Builds a DocumentConverter with proper OCR and table structure options."""
    pipeline_options = PdfPipelineOptions()

    # Enable OCR with specified languages
    pipeline_options.do_ocr = True
    pipeline_options.ocr_options = EasyOcrOptions(lang=languages)

    # For scanned documents, force OCR on every page
    if force_full_page_ocr:
        pipeline_options.ocr_options.force_full_page_ocr = True

    # Use accurate table structure model for best results
    pipeline_options.do_table_structure = True
    pipeline_options.table_structure_options.do_cell_matching = True
    pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )
    return converter


def local_table_reconstruction(
    pdf_input,
    docx_output=None,
    languages=None,
    pages_filter=None,
    auto_open=True,
    force_full_page_ocr=False,
):
    """Processes document structures locally and builds native tables inside Word."""
    if not os.path.exists(pdf_input):
        print(f"❌ Error: File '{pdf_input}' does not exist.")
        sys.exit(1)

    if not docx_output:
        docx_output = os.path.splitext(pdf_input)[0] + "_structured.docx"

    if not languages:
        languages = ["en", "vi"]

    # Get total page count via PyMuPDF
    try:
        pdf_doc = fitz.open(pdf_input)
        total_pages = len(pdf_doc)
        pdf_doc.close()
    except Exception as e:
        print(f"❌ Error reading file structure: {e}")
        sys.exit(1)

    page_range_tuple = parse_pages_to_range_tuple(pages_filter, total_pages)

    print(f"\n📂 Source PDF: {os.path.abspath(pdf_input)}")
    print(f"💾 Target Word: {os.path.abspath(docx_output)}")
    print(f"🌐 OCR Languages: {languages}")
    print(f"🔬 Table Model: ACCURATE (high fidelity)")
    print(f"📷 Full-page OCR: {'Yes (scanned mode)' if force_full_page_ocr else 'No (auto-detect)'}")
    if page_range_tuple:
        print(
            f"📄 Target Page Range: Pages {page_range_tuple[0]} to "
            f"{page_range_tuple[1]} (out of {total_pages} total pages)"
        )
    else:
        print(f"📄 Target Page Range: All Pages (1-{total_pages})")

    print("🔄 Initializing Docling layout analyzer engine...")

    try:
        converter = build_converter(languages, force_full_page_ocr)
        print("🧠 Running document parsing neural networks (extracting text & table frames)...")

        convert_kwargs = {}
        if page_range_tuple:
            convert_kwargs["page_range"] = page_range_tuple

        conversion_result = converter.convert(pdf_input, **convert_kwargs)
        docling_doc = conversion_result.document

        doc = Document()

        print("📈 Reconstructing formatting trees sequentially into Word output...")

        table_count = 0
        text_count = 0

        for item, level in docling_doc.iterate_items():

            if isinstance(item, TableItem):
                table_count += 1
                print(f"   📊 Found table #{table_count}. Mapping cells to Word grid...")
                build_docx_table(doc, item, docling_doc)

            elif isinstance(item, SectionHeaderItem):
                heading_text = item.text.strip()
                if heading_text:
                    # SectionHeaderItem has its own .level attribute
                    heading_level = getattr(item, "level", level)
                    # python-docx heading levels: 0-9 (0 = Title, 1-9 = Heading 1-9)
                    heading_level = max(1, min(heading_level, 9))
                    doc.add_heading(heading_text, level=heading_level)

            elif isinstance(item, ListItem):
                list_text = item.text.strip()
                if list_text:
                    # Use 'List Bullet' style if available, otherwise plain paragraph
                    try:
                        doc.add_paragraph(list_text, style="List Bullet")
                    except KeyError:
                        doc.add_paragraph(f"• {list_text}")

            elif isinstance(item, TextItem):
                text_content = item.text.strip()
                if text_content:
                    doc.add_paragraph(text_content)
                    text_count += 1

            elif isinstance(item, PictureItem):
                # Attempt to extract and insert the picture
                try:
                    pil_image = item.get_image(docling_doc)
                    if pil_image is not None:
                        import io
                        img_buffer = io.BytesIO()
                        pil_image.save(img_buffer, format="PNG")
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(5.0))
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"   🖼️ Inserted picture into document.")
                except Exception as e:
                    # Pictures may not always have extractable image data
                    caption = getattr(item, "caption", None)
                    if caption:
                        doc.add_paragraph(f"[Figure: {caption}]")

        doc.save(docx_output)
        print(
            f"\n🎉 Success! Reconstructed {text_count} text blocks and "
            f"{table_count} tables with fully editable cell text."
        )

        if auto_open:
            open_file_locally(docx_output)

    except Exception as e:
        print(f"\n❌ Layout Analyzer Pipeline Error: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    program_desc = (
        "===============================================================\n"
        "⚡   LOCAL PDF TO WORD CONVERTER WITH TABLE RECONSTRUCTION   ⚡\n"
        "===============================================================\n"
        "Converts complex scanned or flat PDFs into editable Word files offline.\n"
        "Uses local neural nets to detect boundaries and rebuild real Word tables."
    )

    epilog_examples = (
        "💡 USAGE EXAMPLES:\n"
        "  1. Standard conversion (Default English + Vietnamese, all pages):\n"
        "     python3 ocr_table_rebuild_v2.py -i internal_report.pdf\n\n"
        "  2. Convert specific page ranges:\n"
        "     python3 ocr_table_rebuild_v2.py -i data.pdf -p 3-7\n\n"
        "  3. Convert with custom language models (e.g., French + English):\n"
        "     python3 ocr_table_rebuild_v2.py -i invoice.pdf -l fr en -p 1\n\n"
        "  4. Force full-page OCR for scanned documents:\n"
        "     python3 ocr_table_rebuild_v2.py -i scanned.pdf --scanned\n"
    )

    parser = argparse.ArgumentParser(
        description=program_desc,
        epilog=epilog_examples,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument(
        "-i", "--input",
        help="Path to the source PDF document that you want to convert.",
    )
    parser.add_argument(
        "-o", "--output",
        help="Optional: Custom path/filename for the generated Word (.docx) file.",
    )
    parser.add_argument(
        "-l", "--lang",
        nargs="+",
        metavar="CODE",
        help="Optional: Space-separated language codes for OCR (e.g., '-l vi en').",
    )
    parser.add_argument(
        "-p", "--pages",
        metavar="RANGE",
        help="Optional: Page boundary range to parse. Examples: '2-5' or '4'.",
    )
    parser.add_argument(
        "--scanned",
        action="store_true",
        default=False,
        help="Force full-page OCR on every page (use for scanned/image-based PDFs).",
    )
    parser.add_argument(
        "--no-open",
        action="store_true",
        default=False,
        help="Do not automatically open the output file after conversion.",
    )

    args = parser.parse_args()

    if not args.input:
        print("=== Local Layout-Aware PDF Table Reconstructor ===")
        user_in = input("Enter PDF file path: ").strip().strip("\"'")
        user_pages = input(
            "Enter specific page range (e.g., '2-5' or '3', blank for all): "
        ).strip()
        user_pages = user_pages if user_pages else None

        print(
            "\n💡 Language Codes: vi (Vietnamese), en (English), "
            "fr (French), ja (Japanese), zh (Chinese)"
        )
        lang_input = (
            input("Enter language codes space-separated (Leave blank for 'en vi'): ")
            .strip()
            .lower()
        )
        user_langs = lang_input.split() if lang_input else ["en", "vi"]

        scanned_input = (
            input("Is this a scanned/image-based PDF? (y/N): ").strip().lower()
        )
        force_scanned = scanned_input in ("y", "yes")

        if user_in:
            local_table_reconstruction(
                user_in,
                args.output,
                languages=user_langs,
                pages_filter=user_pages,
                force_full_page_ocr=force_scanned,
            )
    else:
        local_table_reconstruction(
            args.input,
            args.output,
            languages=args.lang,
            pages_filter=args.pages,
            auto_open=not args.no_open,
            force_full_page_ocr=args.scanned,
        )
